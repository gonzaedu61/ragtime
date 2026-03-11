import os
import json
from datetime import datetime
from docx import Document
from Utilities import Simple_Progress_Bar


class Docx_Rewriter:
    """
    Structure-preserving DOCX rewriter.
    Rewrites entire paragraphs/headings as single units,
    skips TOC, headers, footers, numeric-only text,
    tracks page numbers using lastRenderedPageBreak,
    logs only text + metadata + result,
    and optionally displays a progress bar.
    """

    def __init__(
        self,
        target_dir: str,
        source_filename: str,
        output_filename: str,
        prompt_template: str,
        llm,
        log_prompts: bool = False,
        progress_enabled: bool = False,
    ):
        self.target_dir = target_dir
        self.source_path = os.path.join(target_dir, source_filename)
        self.output_path = os.path.join(target_dir, output_filename)
        self.prompt_template = prompt_template
        self.llm = llm
        self.log_prompts = log_prompts
        self.progress_enabled = progress_enabled

        if log_prompts:
            self.log_path = os.path.join(
                target_dir,
                f"rewrite_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.jsonl"
            )

        # Progress bar placeholder (initialized later)
        self.progress = None

    # ----------------------------------------------------------------------
    # Detect if a run belongs to a Table of Contents
    # ----------------------------------------------------------------------
    def _is_toc_run(self, run) -> bool:
        r_element = run._r

        for instr in r_element.iter():
            if instr.tag.endswith("instrText") and instr.text and "TOC" in instr.text:
                return True

        for fld in r_element.iter():
            if fld.tag.endswith("fldChar"):
                fld_type = fld.attrib.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fldCharType",
                    ""
                )
                if fld_type in ("begin", "separate", "end"):
                    return True

        return False

    # ----------------------------------------------------------------------
    # Skip numeric-only or non-translatable text
    # ----------------------------------------------------------------------
    def _should_skip_text(self, text: str) -> bool:
        if not text.strip():
            return True

        # Skip if text contains NO alphabetic characters
        if not any(ch.isalpha() for ch in text):
            return True

        return False

    # ----------------------------------------------------------------------
    # Detect page breaks using lastRenderedPageBreak
    # ----------------------------------------------------------------------
    def _paragraph_has_pagebreak(self, paragraph) -> bool:
        p = paragraph._p
        for elem in p.iter():
            if elem.tag.endswith("lastRenderedPageBreak"):
                return True
        return False

    # ----------------------------------------------------------------------
    # Logging helper
    # ----------------------------------------------------------------------
    def _log(self, entry: dict):
        if not self.log_prompts:
            return
        with open(self.log_path, "a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")

    # ----------------------------------------------------------------------
    # Rewrite a full paragraph
    # ----------------------------------------------------------------------
    def _rewrite_paragraph(self, paragraph, metadata):
        full_text = paragraph.text.strip()

        if self._should_skip_text(full_text):
            return

        rewritten = self._rewrite_text(full_text, metadata)

        # Clear existing runs
        for run in paragraph.runs:
            run.text = ""

        # Insert rewritten text as a single run
        paragraph.add_run(rewritten)

    # ----------------------------------------------------------------------
    # Rewrite a single text unit
    # ----------------------------------------------------------------------
    def _rewrite_text(self, text: str, metadata: dict) -> str:
        prompt = self.prompt_template.format(text=text)

        response = self.llm.complete(prompt)
        #response = text  # placeholder for now

        # Log only text + metadata + result
        self._log({
            "metadata": metadata,
            "text": text,
            "response": response,
        })

        # Progress bar update
        if self.progress:
            self.progress.update(step=1, label=f"Page {metadata.get('page')}")

        return response

    # ----------------------------------------------------------------------
    # Count total paragraphs to rewrite (for progress bar)
    # ----------------------------------------------------------------------
    def _count_total_paragraphs(self, doc):
        total = 0

        # Main body
        total += len(doc.paragraphs)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total += len(cell.paragraphs)

        # Headers & footers
        for section in doc.sections:
            total += len(section.header.paragraphs)
            total += len(section.footer.paragraphs)

        return total

    # ----------------------------------------------------------------------
    # Main processing
    # ----------------------------------------------------------------------
    def rewrite(self):
        doc = Document(self.source_path)

        # Initialize progress bar
        if self.progress_enabled:
            total = self._count_total_paragraphs(doc)
            self.progress = Simple_Progress_Bar(total=total, enabled=True)
        else:
            self.progress = None

        current_page = 1

        # -------------------------
        # Rewrite main body
        # -------------------------
        for p_idx, paragraph in enumerate(doc.paragraphs):

            if self._paragraph_has_pagebreak(paragraph):
                current_page += 1

            if any(self._is_toc_run(run) for run in paragraph.runs):
                continue

            metadata = {
                "type": "paragraph",
                "paragraph_index": p_idx,
                "page": current_page,
            }

            self._rewrite_paragraph(paragraph, metadata)

        # -------------------------
        # Rewrite tables
        # -------------------------
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p_idx, paragraph in enumerate(cell.paragraphs):

                        if self._paragraph_has_pagebreak(paragraph):
                            current_page += 1

                        if any(self._is_toc_run(run) for run in paragraph.runs):
                            continue

                        metadata = {
                            "type": "table_cell",
                            "table_index": t_idx,
                            "row_index": r_idx,
                            "col_index": c_idx,
                            "paragraph_index": p_idx,
                            "page": current_page,
                        }

                        self._rewrite_paragraph(paragraph, metadata)

        # -------------------------
        # Rewrite headers & footers
        # -------------------------
        for s_idx, section in enumerate(doc.sections):

            # ---- Headers ----
            header = section.header
            for p_idx, paragraph in enumerate(header.paragraphs):

                metadata = {
                    "type": "header",
                    "section_index": s_idx,
                    "paragraph_index": p_idx,
                    "page": None,
                }

                if any(self._is_toc_run(run) for run in paragraph.runs):
                    continue

                self._rewrite_paragraph(paragraph, metadata)

            # ---- Footers ----
            footer = section.footer
            for p_idx, paragraph in enumerate(footer.paragraphs):

                metadata = {
                    "type": "footer",
                    "section_index": s_idx,
                    "paragraph_index": p_idx,
                    "page": None,
                }

                if any(self._is_toc_run(run) for run in paragraph.runs):
                    continue

                self._rewrite_paragraph(paragraph, metadata)

        doc.save(self.output_path)
        return self.output_path
